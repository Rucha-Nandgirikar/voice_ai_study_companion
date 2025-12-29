from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path

import httpx
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from dotenv import load_dotenv

from backend.schemas import (
    ExtractRequest,
    ExtractResponse,
    ExtractPartsRequest,
    ExtractPartsResponse,
    ExtractPart,
    TopicsRequest,
    TopicsResponse,
    TopicItem,
    NotesAppendQuestionRequest,
    NotesAppendTurnRequest,
    NotesAppendQARequest,
    NotesAppendQuizRequest,
    NotesGetResponse,
    NotesResetRequest,
    NotesSetSummaryRequest,
    SessionStartRequest,
    SessionLatestResponse,
    SessionItem,
    SessionsListResponse,
    ProgressMarkTopicDoneRequest,
    ProgressResponse,
)
from backend.url_extract import fetch_and_extract_main_text, fetch_and_extract_topics
from backend.study_repo import PostgresStudyRepo, make_study_repo


# Local dev convenience:
# - If you run from repo root, `.env` is loaded.
# - If you run from inside `backend/` (or other CWD), `backend/.env` is also loaded.
# Cloud Run should use real env vars / Secret Manager (it will not have your local .env file).
load_dotenv()
load_dotenv(dotenv_path=Path(__file__).with_name(".env"), override=False)

openapi_tags = [
    {"name": "System", "description": "Health checks and service info"},
    {"name": "Sessions", "description": "Session sidebar list and housekeeping"},
    {"name": "Topics", "description": "Topic discovery + selected topics for today"},
    {"name": "Agent Tools", "description": "Tools intended to be called by the ElevenLabs agent"},
    {"name": "Notes", "description": "Persisted notes + DOCX export"},
]

app = FastAPI(
    title="Voice AI Study Companion API",
    version="0.2.0",
    openapi_tags=openapi_tags,
)
repo = make_study_repo()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/", tags=["System"])
def root() -> dict:
    return {
        "ok": True,
        "service": "voice-ai-study-companion",
        "endpoints": [
            "/health",
            "/extract",
            "/topics",
            "/sessions",
            "/sessions/start",
            "/sessions/latest",
            "/sessions (DELETE)",
            "/notes/reset",
            "/notes/set_summary",
            "/notes/append_question",
            "/notes/append_turn",
            "/notes/append_qa",
            "/notes/append_quiz",
            "/notes",
            "/notes/download.docx",
        ],
        "docs": "/docs",
    }


@app.get("/health", tags=["System"])
def health() -> dict:
    return {"ok": True}


@app.on_event("startup")
def _startup() -> None:
    if isinstance(repo, PostgresStudyRepo):
        repo.ensure_schema()


@app.get("/sessions", response_model=SessionsListResponse, tags=["Sessions"])
def sessions_list(limit: int = 50) -> SessionsListResponse:
    try:
        sessions = repo.list_sessions(limit=limit)
        items = [
            SessionItem(
                sessionId=s.session_id,
                url=s.url,
                selectedTopics=s.selected_topics,
                completedTopics=s.completed_topics,
                currentTopic="",  # Deprecated: no longer used
                createdAt=s.created_at,
                updatedAt=s.updated_at,
            )
            for s in sessions
        ]
        return SessionsListResponse(sessions=items)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Sessions list failed: {e}")


@app.post("/sessions/start", response_model=SessionItem, tags=["Sessions"])
def sessions_start(req: SessionStartRequest) -> SessionItem:
    try:
        s = repo.start_session(req.url, req.selectedTopics or [])
        return SessionItem(
            sessionId=s.session_id,
            url=s.url,
            selectedTopics=s.selected_topics,
            completedTopics=s.completed_topics,
            currentTopic="",  # Deprecated: no longer used
            createdAt=s.created_at,
            updatedAt=s.updated_at,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Sessions start failed: {e}")


@app.get("/sessions/latest", response_model=SessionLatestResponse, tags=["Sessions"])
def sessions_latest(url: str) -> SessionLatestResponse:
    try:
        s = repo.latest_session_for_url(url)
        if not s:
            return SessionLatestResponse(session=None)
        return SessionLatestResponse(
            session=SessionItem(
                sessionId=s.session_id,
                url=s.url,
                selectedTopics=s.selected_topics,
                completedTopics=s.completed_topics,
                currentTopic="",  # Deprecated: no longer used
                createdAt=s.created_at,
                updatedAt=s.updated_at,
            )
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Sessions latest failed: {e}")


@app.delete("/sessions", tags=["Sessions"])
def sessions_delete(sessionId: str) -> dict:
    try:
        repo.delete_session(sessionId)
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Sessions delete failed: {e}")


@app.post("/sessions/progress/mark_topic_done", response_model=ProgressResponse, tags=["Agent Tools"])
def progress_mark_topic_done(req: ProgressMarkTopicDoneRequest) -> ProgressResponse:
    """
    Mark a topic as completed for this session.
    The topic must exist in selectedTopics for this session.
    """
    try:
        s = repo.mark_topic_done(req.sessionId, req.topicTitle)
        return ProgressResponse(
            sessionId=s.session_id,
            url=s.url,
            selectedTopics=s.selected_topics,
            completedTopics=s.completed_topics,
            currentTopic="",  # Deprecated: no longer used
            updatedAt=s.updated_at,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Progress mark topic done failed: {e}")


@app.post("/extract", response_model=ExtractResponse, tags=["Agent Tools"])
async def extract(req: ExtractRequest) -> ExtractResponse:
    """
    Option B backend: the agent handles all LLM calls (Gemini configured in ElevenLabs).
    This endpoint only fetches & extracts the main page text.
    """
    try:
        text = await fetch_and_extract_main_text(req.url)
        if not text or len(text) < 200:
            raise HTTPException(status_code=400, detail="Could not extract enough readable text from that URL.")
        total = len(text)
        # Prevent oversized tool payloads (ElevenLabs tool results/context can have size limits).
        default_max = int(os.environ.get("EXTRACT_MAX_CHARS", "12000"))
        max_chars = int(req.maxChars) if req.maxChars is not None else default_max
        max_chars = max(500, min(max_chars, 200_000))
        if total > max_chars:
            clipped = text[:max_chars].rstrip()
            clipped += "\n\n[TRUNCATED]\nAsk for more by calling /extract/chunk with a higher offset."
            return ExtractResponse(url=req.url, cleanedText=clipped, truncated=True, totalChars=total)
        return ExtractResponse(url=req.url, cleanedText=text, truncated=False, totalChars=total)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Extract failed: {e}")


def _split_into_parts(text: str, num_parts: int, max_chars_per_part: int) -> tuple[list[tuple[int, str]], bool]:
    """
    Split text into `num_parts` parts, aiming for paragraph boundaries.
    Returns (parts, truncated) where parts is list of (offset, part_text).
    Always returns exactly num_parts parts; some may be empty if input is short.
    """
    if num_parts < 2:
        num_parts = 2
    max_chars_per_part = max(500, min(int(max_chars_per_part), 200_000))

    total = len(text)
    if total == 0:
        return ([(0, "") for _ in range(num_parts)], False)

    # If the text is huge, only deliver up to num_parts * max_chars_per_part via this endpoint.
    deliver_cap = num_parts * max_chars_per_part
    truncated = total > deliver_cap
    deliver_text = text[:deliver_cap] if truncated else text

    paras = [p.strip() for p in deliver_text.split("\n\n") if p.strip()]
    if not paras:
        # Fallback: hard split by char length
        parts: list[tuple[int, str]] = []
        for i in range(num_parts):
            start = i * max_chars_per_part
            parts.append((start, deliver_text[start : start + max_chars_per_part]))
        return (parts, truncated)

    # Target size per part based on deliver_text length.
    target = max(1, len(deliver_text) // num_parts)
    parts_text: list[str] = []
    current: list[str] = []
    current_len = 0

    for p in paras:
        p_len = len(p) + (2 if current else 0)
        if current and (current_len + p_len > target) and (len(parts_text) < num_parts - 1):
            parts_text.append("\n\n".join(current).strip())
            current = [p]
            current_len = len(p)
        else:
            current.append(p)
            current_len += p_len

    if current:
        parts_text.append("\n\n".join(current).strip())

    # Normalize to exactly num_parts parts.
    if len(parts_text) < num_parts:
        parts_text.extend([""] * (num_parts - len(parts_text)))
    elif len(parts_text) > num_parts:
        # Merge overflow into the last part.
        head = parts_text[: num_parts - 1]
        tail = "\n\n".join(parts_text[num_parts - 1 :]).strip()
        parts_text = head + [tail]

    # Enforce max_chars_per_part for each part (hard cap).
    capped: list[str] = []
    for pt in parts_text:
        if len(pt) > max_chars_per_part:
            capped.append(pt[:max_chars_per_part].rstrip() + "\n\n[PART TRUNCATED]")
        else:
            capped.append(pt)

    # Compute offsets by searching sequentially (best-effort).
    offsets: list[int] = []
    cursor = 0
    for pt in capped:
        if not pt:
            offsets.append(cursor)
            continue
        idx = deliver_text.find(pt.replace("\n\n[PART TRUNCATED]", "").rstrip(), cursor)
        if idx == -1:
            offsets.append(cursor)
        else:
            offsets.append(idx)
            cursor = idx + len(pt)

    return (list(zip(offsets, capped)), truncated)


@app.post("/extract/parts", response_model=ExtractPartsResponse, tags=["Agent Tools"], include_in_schema=False)
async def extract_parts(req: ExtractPartsRequest) -> ExtractPartsResponse:
    """
    Fetches & extracts the page text, then returns it split into N parts (default 4).
    This is designed for agents/LLMs to process large pages safely by summarizing each part.
    """
    try:
        text = await fetch_and_extract_main_text(req.url)
        if not text or len(text) < 200:
            raise HTTPException(status_code=400, detail="Could not extract enough readable text from that URL.")

        num_parts = int(req.parts or 4)
        max_chars_per_part = int(req.maxCharsPerPart or 9000)
        parts_with_offsets, truncated = _split_into_parts(text, num_parts=num_parts, max_chars_per_part=max_chars_per_part)
        parts = [
            ExtractPart(index=i + 1, totalParts=num_parts, offset=off, text=pt)
            for i, (off, pt) in enumerate(parts_with_offsets)
        ]
        return ExtractPartsResponse(url=req.url, title=None, parts=parts, truncated=truncated, totalChars=len(text))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Extract parts failed: {e}")


@app.post("/topics", response_model=TopicsResponse, tags=["Topics"])
async def topics(req: TopicsRequest) -> TopicsResponse:
    """
    Best-effort topic listing from the page structure (headings / markdown headings).
    This endpoint does NOT call an LLM; it is cheap and fast.
    """
    try:
        title, topics_pairs = await fetch_and_extract_topics(req.url)
        topics_items = [TopicItem(level=lvl, title=t) for (lvl, t) in topics_pairs]
        return TopicsResponse(url=req.url, title=title, topics=topics_items)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Topics failed: {e}")


@app.post("/topics/select", tags=["Topics"], include_in_schema=False)
def topics_select_legacy(_: dict) -> dict:
    raise HTTPException(status_code=410, detail="Deprecated. Use POST /sessions/start to store selected topics.")


@app.get("/topics/selected", tags=["Topics"], include_in_schema=False)
def topics_selected_legacy(_: str) -> dict:
    raise HTTPException(status_code=410, detail="Deprecated. Use GET /sessions/latest?url=... to read selected topics.")


@app.post("/notes/reset", response_model=NotesGetResponse, tags=["Notes"])
def notes_reset(req: NotesResetRequest) -> NotesGetResponse:
    try:
        rec = repo.reset_notes(req.sessionId)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes reset failed: {e}")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.post("/notes/set_summary", response_model=NotesGetResponse, tags=["Notes"])
def notes_set_summary(req: NotesSetSummaryRequest) -> NotesGetResponse:
    try:
        rec = repo.set_summary(req.sessionId, req.summary)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes set_summary failed: {e}")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.post("/notes/append_question", response_model=NotesGetResponse, tags=["Notes"], include_in_schema=False)
def notes_append_question(req: NotesAppendQuestionRequest) -> NotesGetResponse:
    try:
        rec = repo.append_turn(req.sessionId, "user", req.question)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes append_question failed: {e}")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.post("/notes/append_turn", response_model=NotesGetResponse, tags=["Notes"])
def notes_append_turn(req: NotesAppendTurnRequest) -> NotesGetResponse:
    try:
        rec = repo.append_turn(req.sessionId, req.role, req.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes append_turn failed: {e}")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.post("/notes/append_qa", response_model=NotesGetResponse, tags=["Notes"])
def notes_append_qa(req: NotesAppendQARequest) -> NotesGetResponse:
    try:
        rec = repo.append_qa(req.sessionId, req.question, req.answer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes append_qa failed: {e}")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.post("/notes/append_quiz", response_model=NotesGetResponse, tags=["Notes"])
def notes_append_quiz(req: NotesAppendQuizRequest) -> NotesGetResponse:
    try:
        rec = repo.append_quiz(req.sessionId, req.question, req.userAnswer, req.correctAnswer, req.explanation)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes append_quiz failed: {e}")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.get("/notes", response_model=NotesGetResponse, tags=["Notes"])
def notes_get(sessionId: str) -> NotesGetResponse:
    try:
        rec = repo.get_notes(sessionId)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notes get failed: {e}")
    if not rec:
        raise HTTPException(status_code=404, detail="Unknown sessionId")
    return NotesGetResponse(
        sessionId=rec.session_id,
        url=rec.url,
        summary=rec.summary,
        questions=rec.questions,
        turns=rec.turns,
        qa=rec.qa,
        quizzes=rec.quizzes,
        updatedAt=rec.updated_at,
    )


@app.get("/notes/download.docx", tags=["Notes"])
def notes_download_docx(sessionId: str) -> StreamingResponse:
    rec = repo.get_notes(sessionId)
    if not rec:
        raise HTTPException(status_code=404, detail="Unknown sessionId")

    doc = Document()
    doc.add_heading("Voice AI Study Notes", level=1)
    doc.add_paragraph(f"Source URL: {rec.url}")
    doc.add_paragraph(f"Last updated: {rec.updated_at}")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(rec.summary or "(No summary saved yet)")

    doc.add_heading("Q&A", level=2)
    if rec.qa:
        for idx, pair in enumerate(rec.qa, start=1):
            q = (pair.get("q") or "").strip()
            a = (pair.get("a") or "").strip()
            if q:
                doc.add_paragraph(f"Q{idx}. {q}")
            if a:
                doc.add_paragraph(f"A{idx}. {a}")
            doc.add_paragraph("")  # spacer
    else:
        doc.add_paragraph("(No Q&A saved yet)")

    doc.add_heading("Quizzes", level=2)
    if rec.quizzes:
        for idx, qz in enumerate(rec.quizzes, start=1):
            q = (qz.get("question") or "").strip()
            ua = (qz.get("userAnswer") or "").strip()
            ca = (qz.get("correctAnswer") or "").strip()
            ex = (qz.get("explanation") or "").strip()
            if q:
                doc.add_paragraph(f"Quiz {idx}: {q}")
            if ua:
                doc.add_paragraph(f"Your answer: {ua}")
            if ca:
                doc.add_paragraph(f"Correct answer: {ca}")
            if ex:
                doc.add_paragraph(f"Explanation: {ex}")
            doc.add_paragraph("")  # spacer
    else:
        doc.add_paragraph("(No quizzes saved yet)")

    # Back-compat sections (optional)
    if rec.turns:
        doc.add_heading("Call transcript (raw)", level=2)
        for t in rec.turns:
            role = (t.get("role") or "").strip().lower()
            text = (t.get("text") or "").strip()
            if not text:
                continue
            prefix = "You:" if role == "user" else "Tutor:"
            doc.add_paragraph(f"{prefix} {text}")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = "study-notes.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )



